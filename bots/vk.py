import asyncio
from services.config_vk import bot
from vkbottle import BaseStateGroup
from vkbottle.bot import Message, MessageEvent
import re
from services.Datbase import DataBase
from typing import Optional, List, Dict, Tuple
from docx import Document
from vkbottle import Keyboard, KeyboardButtonColor, Callback, GroupEventType
from math import ceil
from random import randint
from dotenv import load_dotenv
from os import getenv
from os.path import exists
import json
from datetime import datetime
from services.logging_config import setup_logging
import logging
from services.cmd_handler import get_groq_response
from ast import literal_eval
import pickle
import aiofiles
from pathlib import Path
import time
from vkbottle import PhotoMessageUploader

uploader = PhotoMessageUploader(bot.api)

data = DataBase()
load_dotenv()
setup_logging()

logger = logging.getLogger('vk_bot')

#const
ITEMS_PER_PAGE = 6
GROUP_ID = getenv('schedule_id')
ASSETS_DIR = Path('assets')

if not data.get_admins():
    vk_id = getenv("owner_vk_id")
    data.ensure_user(vk_id=vk_id,role="admin",full_name="Вячеслав М")

def convert_group_name(group_name:str = None) -> str:
    if group_name is None:
        return None
    
    result = re.sub(r'\D', '', group_name)
    return result

class ServeyState(BaseStateGroup):
    JOIN = "join"

class SetterState:
    def __init__(self, parent):
        self.parent = parent
    
    async def lastCmid(self, peer_id:Optional[int] = None, 
                          cmid:Optional[int] = None) -> None:
        if (peer_id is None) or (cmid is None):
            raise ValueError("peer_id is None" if peer_id is None else "cmid is None")
        
        await self.parent.load()
        
        if self.parent.GROUPS.get(peer_id) is None:
            await self.parent.add(peer_id=peer_id)
            await self.parent.load()
        
        lucmid = self.parent.GROUPS.get(peer_id, {}).get("last_user_cmid") or 1

        if (cmid - lucmid) > 1:
            lucmid = lucmid if (cmid - lucmid) < 100 else (cmid - 100)
            for i in range(lucmid, cmid+1):
                try:
                    response = await bot.api.messages.get_by_conversation_message_id(
                        conversation_message_ids=[i],
                        peer_id=peer_id
                    )
                    message = response.items[0] if response.items else None
                    from_id = getattr(message, "from_id", None) if message != None else None
                    if from_id < 0:
                        self.parent.GROUPS[peer_id]["last_bot_cmid"] = i
                except Exception:
                    logger.exception("Ошибка при проверке cmid")

        
        await self.parent.save()
    
    async def lastSchCmid(self, peer_id:Optional[int] = None):
        if peer_id is None:
            raise ValueError("peer_id is None")
        
        if self.parent.GROUPS.get(peer_id) is None:
            await self.parent.add(peer_id)
            await self.parent.load()
        
        bot_cmid = self.parent.GROUPS.get(peer_id, {}).get("last_bot_cmid")
        self.parent.GROUPS[peer_id]["last_sch_cmid"] = bot_cmid
        
        


class GetterState:
    def __init__(self, parent):
        self.parent = parent
    
    async def userCmid(self, peer_id:Optional[int] = None) -> int:
        await self.parent.load()
        
        if self.parent.GROUPS.get(peer_id) is None:
            await self.parent.add(peer_id)
            await self.parent.load()

        cmid = self.parent.GROUPS.get(peer_id, {}).get("last_user_cmid") or 1
        return cmid
    
    async def botCmid(self, peer_id:Optional[int] = None) -> int:
        await self.parent.load()
        
        if self.parent.GROUPS.get(peer_id) is None:
            await self.parent.add(peer_id)
            await self.parent.load()

        cmid = self.parent.GROUPS.get(peer_id, {}).get("last_bot_cmid") or None
        return cmid
    
    async def schCmid(self, peer_id:Optional[int] = None) -> int:
        await self.parent.load()
        
        if self.parent.GROUPS.get(peer_id) is None:
            await self.parent.add(peer_id)
            await self.parent.load()

        cmid = self.parent.GROUPS.get(peer_id, {}).get("last_sch_cmid") or None
        return cmid
        
        


class StateCmid:
    def __init__(self):
        self.filename = "stateCmid.pkl"
        self.GROUPS = {}
        self.set = SetterState(self)
        self.get = GetterState(self)
    
    @classmethod
    async def create(cls) -> StateCmid:
        instanse = cls()
        await instanse.load()
        return instanse
    
    async def load(self) -> None:
        if exists(self.filename):
            try:
                with open(self.filename, 'rb') as f:
                    self.GROUPS = pickle.load(f)
                logger.info("Данные состояния загружены")
            except Exception:
                logger.exception("Ошибка при загрузке состояния: ")
                self.GROUPS = {}
        else:
            logger.info("Файл состояния не найден")
    
    async def save(self,) -> None:
        try:
            with open(self.filename, "wb") as f:
                pickle.dump(self.GROUPS, f)
            logger.info("Состояние сохранено!")
        except Exception:
            logger.exception("Ошибка при сохранении состояний: ")
    
    async def add(self, peer_id:Optional[int] = None) -> None:
        if peer_id is None:
            raise ValueError("peer_id is None")
        await self.load()
        self.GROUPS[peer_id] = {
            "last_user_cmid":1,
            "last_bot_cmid":1,
            "last_sch_cmid":1
        }
        await self.save()
    



stateCmid = StateCmid()


def get_groups() -> List:
    doc = Document((ASSETS_DIR / "schedule.docx"))
    groups = []

    if not doc.tables:
        return None
    
    table = doc.tables[0]
    data = DataBase()

    for row in table.rows:
        cells = clean([cell.text.strip() for cell in row.cells])

        if len(cells) == 4:
                groups.append(convert_group_name(cells[0]))
        elif len(cells) > 4:
            groups.append(convert_group_name(cells[1]))
    
    return clean(groups)


def clean(items:List) -> List:
    result = []
    for item in items:
        if item and item not in result:
            result.append(item)
    return result

def join_kb(page:int) -> Keyboard:
    kb = Keyboard(inline=True)

    names = get_groups()

    start_idx = page * ITEMS_PER_PAGE
    end_idx = start_idx + ITEMS_PER_PAGE
    length = len(names)

    if max([start_idx, end_idx, length-1]) != (length-1):
        return None
    
    cur_names = names[start_idx:end_idx]

    for i, name in enumerate(cur_names):
        kb.add(
            Callback(name, payload={"act":"select_name", "name":name}),
            color=KeyboardButtonColor.PRIMARY
        )

        if (i + 1) % 2 == 0 and (i + 1) != len(cur_names):
            kb.row()
        
    kb.row()

    total_page = ceil(length/ITEMS_PER_PAGE)
    if page > 0:
        kb.add(
            Callback("<- Назад", payload={"act":"change_page", "page":page-1,"method":"join_kb"}),
            color=KeyboardButtonColor.SECONDARY
        )
    
    if (page + 1) < total_page:
        kb.add(
            Callback("Вперёд ->", payload={"act":"change_page", "page":page+1, "method":"join_kb"}),
            color=KeyboardButtonColor.SECONDARY
        )
    
    kb.row()

    kb.add(
        Callback("Закрыть", payload={"act":"close"}),
        color=KeyboardButtonColor.NEGATIVE
    )
    
    return kb

def hw_subj_kb(vk_id:int = None, page:int = 0) -> Keyboard:
    global ITEMS_PER_PAGE

    if vk_id is None:
        return

    kb = Keyboard(inline=True)

    subjects = data.get_subjects(vk_id=vk_id)

    start_idx = page * ITEMS_PER_PAGE
    end_idx = start_idx + ITEMS_PER_PAGE
    length = len(subjects)

    if max([start_idx, end_idx, length-1]) != (length-1):
        return
    
    cur_subjects = subjects[start_idx:end_idx]

    for i,(id,name) in enumerate(cur_subjects):
        kb.add(
            Callback(f"{name}", payload={"act":"get_hw", "subj_id":id}),
            KeyboardButtonColor.PRIMARY
        )

        if ((i+1) % 2 == 0) and (i + 1 != len(cur_subjects)):
            kb.row()
    
    kb.row()

    total_page = ceil(length/ITEMS_PER_PAGE)
    if page > 0:
        kb.add(
            Callback("<- Назад", payload={"act":"change_page", "page":page-1,"method":"hw_subj_kb"}),
            color=KeyboardButtonColor.SECONDARY
        )
    
    if (page + 1) < total_page:
        kb.add(
            Callback("Вперёд ->", payload={"act":"change_page", "page":page+1, "method":"hw_subj_kb"}),
            color=KeyboardButtonColor.SECONDARY
        )
    
    kb.row()

    kb.add(
        Callback("Закрыть", payload={"act":"close"}),
        color=KeyboardButtonColor.NEGATIVE
    )
    
    return kb

def close_kb():
    kb = Keyboard(inline=True)
    kb.add(
        Callback("Закрыть",payload={"act":"close"}),
        color=KeyboardButtonColor.NEGATIVE
    )
    return kb

def remove_kb(vk_id:int, page:int = 0) -> Keyboard:
    global ITEMS_PER_PAGE

    if vk_id is None:
        return

    kb = Keyboard(inline=True)

    subjects = data.get_subjects(vk_id=vk_id)

    start_idx = page * ITEMS_PER_PAGE
    end_idx = start_idx + ITEMS_PER_PAGE
    length = len(subjects)

    if max([start_idx, end_idx, length-1]) != (length-1):
        return
    
    cur_subjects = subjects[start_idx:end_idx]

    for i,(id,name) in enumerate(cur_subjects):
        kb.add(
            Callback(f"{name}", payload={"act":"del_hw", "subj_id":id}),
            KeyboardButtonColor.PRIMARY
        )

        if ((i+1) % 2 == 0) and (i + 1 != len(cur_subjects)):
            kb.row()
    
    kb.row()

    total_page = ceil(length/ITEMS_PER_PAGE)
    if page > 0:
        kb.add(
            Callback("<- Назад", payload={"act":"change_page", "page":page-1,"method":"hw_subj_kb"}),
            color=KeyboardButtonColor.SECONDARY
        )
    
    if (page + 1) < total_page:
        kb.add(
            Callback("Вперёд ->", payload={"act":"change_page", "page":page+1, "method":"hw_subj_kb"}),
            color=KeyboardButtonColor.SECONDARY
        )
    
    kb.row()

    kb.add(
        Callback("Закрыть", payload={"act":"close"}),
        color=KeyboardButtonColor.NEGATIVE
    )
    
    return kb

def confirm_adm_kb(name:str, peer_id:int) -> Keyboard:
    kb = Keyboard(inline=True)

    kb.add(
        Callback("Нет", payload={"act":"rejected","name":name, "peer_id":peer_id}),
        color=KeyboardButtonColor.NEGATIVE
    )

    kb.add(
        Callback("Да", payload={"act":"accepted","name":name, "peer_id":peer_id}),
        color=KeyboardButtonColor.POSITIVE
    )

    return kb

def confirm_homework_kb(message_text:str = None, cmid:int = None) -> Keyboard:
    if message_text is None or cmid is None:
        return None

    kb = Keyboard(inline=True)

    kb.add(
        Callback("Да", payload={"act":"is_homework","text":message_text,"cmid":cmid}),
        color=KeyboardButtonColor.POSITIVE
    )

    kb.add(Callback("Нет", payload={"act":"!is_homework"}),
           color=KeyboardButtonColor.NEGATIVE)
    
    return kb


@bot.on.chat_message(func=lambda m: bool(m.text and re.search(r'\bбот\b', m.text.lower())))
async def start_message(message: Message):

    logger.info("Бот получил команду: '%s' от пользователя %s в группе %s", message.text,
                 message.from_id,message.peer_id)
    cur_state = await bot.state_dispenser.get(message.peer_id)
    if cur_state:
        cur_state = cur_state.state
    else:
        cur_state = ""
    text = message.text.lower()
    
    if (text.strip().endswith("привяжи")) and (cur_state != ServeyState.JOIN):
        keyboard = join_kb(0)
        await bot.api.messages.send(peer_id=message.peer_id,
                                    random_id=randint(0,10000),
                                    keyboard=keyboard,
                                    message="Выберите группу",
                                    silent=True)
        await bot.api.messages.delete(peer_id=message.peer_id,
                                      cmids=[message.conversation_message_id],
                                      delete_for_all=True)
        await bot.state_dispenser.set(message.peer_id, ServeyState.JOIN)
    elif (text.strip().endswith("обнови базу")):
        response = await bot.api.messages.get_conversation_members(
            peer_id=message.peer_id
        )
        group_id = data.get_group_id(vk_id=message.peer_id)
        if group_id is None:
            return

        users = {user.id:f"{user.first_name}{user.last_name}" for user in response.profiles}
        for member in response.items:
            member_id = member.member_id

            if (member_id > 0) and (member_id in users):
                name = users[member_id]
                if data.ensure_user(full_name=f"{name}", vk_id=member_id, group_id=group_id) is not None:
                    logger.info(f"Новый пользователь {name}|{member_id} привязан к группе {group_id}!")
    elif (text.strip().endswith("дз")):
        keyboard = hw_subj_kb(vk_id=message.peer_id, page = 0)
        await bot.api.messages.send(peer_id=message.peer_id,
                                    random_id=randint(0,10000),
                                    keyboard=keyboard,
                                    message="Выберите предмет",
                                    silent=True)
        await bot.api.messages.delete(peer_id=message.peer_id,
                                      cmids=[message.conversation_message_id],
                                      delete_for_all=True)
    elif (text.strip().endswith("расписание")):
        await bot.api.messages.delete(peer_id=message.peer_id,
                                      cmids=[message.conversation_message_id],
                                      delete_for_all=True)
        name = data.get_group_name(vk_id=message.peer_id)
        schedule = data.get_schedule(group_name=name)
        user_id = message.from_id
        users_info = await bot.api.users.get(user_ids=[user_id], fields=["screen_name"])
        if users_info:
            user = users_info[0]
            # Если у пользователя установлен короткий адрес (domain/screen_name)
            if user.screen_name:
                nick =  f"@{user.screen_name}"
            else:
                # Если ника нет, возвращаем ссылку через id
                nick =  f"@id{user_id}"
        else: nick = "Неизвестно"
        lines = [f"Расписание, {nick}:"]
        ln = []
        subjects = []
        
        for lesson, subject, room in schedule:
            l = re.sub(r'\D', '', lesson)
            ln.append(l)
            s = subject.capitalize()
            subjects.append(s)
            r = "".join(c for c in room if c.isdigit()) if any(c.isdigit() for c in room) else room
            lines.append(f"({l})  {s}  [{r}]")
        
        text = "\n".join(lines)
        keyboard = close_kb()
        await bot.api.messages.send(peer_id=message.peer_id, random_id = randint(0,10000),
                                    silent = True, keyboard=keyboard, message=text)
        
    elif (text.strip().endswith("удали")):
        kb = remove_kb(message.peer_id)
        await bot.api.messages.send(peer_id=message.peer_id, random_id=randint(0,1000),
                                    silent=True, message = "Выберите предмет:", keyboard = kb)
        await bot.api.messages.delete(cmids=[message.conversation_message_id],peer_id=message.peer_id,
                                      delete_for_all=True)
        

    else:
        all_chunks = await get_all_texts_recursive(message)
        full_text = ". ".join(all_chunks)
        result = await get_groq_response(message=full_text,group_name=data.get_group_name(vk_id=message.peer_id))

        if result != "None":
            result = literal_eval(result.strip())
            if len(result) != 3:
                return
            
            att_urls = await get_all_photos_recursive(message)
            att_names = []
            for url in att_urls:
                name = await download_photo(url)
                att_names.append(name)

            
            subject_name = data.get_subject_name(subject_id=result[0])
            data.ensure_homework(vk_id=message.peer_id, subject_name=subject_name, description=f"{result[1]}", attachments=att_names)
            keyboard = close_kb()
            await bot.api.messages.send(peer_id=message.peer_id,
                                        random_id=randint(0,100000),
                                        silent = True,
                                        message=f"ДЗ по {subject_name} добавлено!",
                                        keyboard=keyboard)


        else:
            for Id in data.get_admins():
                vk_id = data.get_user_vk_id(Id)
                await bot.api.messages.send(
                    user_id=vk_id,
                    message=f"Не получилось добавить дз '{text}' в группе {message.peer_id} пользователем {message.from_id}",
                    random_id=randint(0,1000000)
                )

async def download_photo(url):
    try:
        content = await bot.api.http_client.request_content(url)
        now = int(time.time())
        file_name = f"photo_{now}.jpeg"
        file_path = ASSETS_DIR / file_name
        async with aiofiles.open(file_path, mode="wb") as f:
            await f.write(content)
        await asyncio.sleep(1)
        return file_name
    except Exception:
        logger.exception("Ошибка при загрузке фото: ")

async def get_all_texts_recursive(msg) -> list:
    """Рекурсивно собирает текст из самого сообщения, ответов и пересланных сообщений."""
    texts = []
    
    # 1. Забираем текст текущего сообщения (если он есть)
    if msg.text:
        texts.append(msg.text.strip())

    # 2. Проверяем ответ на сообщение (reply_message)
    # В vkbottle это объект, а не список
    if getattr(msg, "reply_message", None):
        texts.extend(await get_all_texts_recursive(msg.reply_message))

    # 3. Проверяем пересланные сообщения (fwd_messages)
    # Это всегда список
    if getattr(msg, "fwd_messages", None):
        for fwd in msg.fwd_messages:
            texts.extend(await get_all_texts_recursive(fwd))
            
    return texts

async def get_all_photos_recursive(msg) -> list:
    """Рекурсивно собирает URL всех фотографий из сообщения и пересылок."""
    photo_urls = []

    # 1. Ищем фото в текущем сообщении
    if getattr(msg, "attachments", None):
        for attachment in msg.attachments:
            if attachment.photo:
                # Берем последний элемент в списке sizes (самый большой)
                max_size_url = attachment.photo.sizes[-1].url
                photo_urls.append(max_size_url)

    # 2. Идем в ответ (reply_message)
    reply = getattr(msg, "reply_message", None)
    if reply:
        photo_urls.extend(await get_all_photos_recursive(reply))

    # 3. Идем в пересланные сообщения (fwd_messages)
    fwd_messages = getattr(msg, "fwd_messages", None)
    if fwd_messages:
        for fwd in fwd_messages:
            photo_urls.extend(await get_all_photos_recursive(fwd))

    return photo_urls

@bot.on.raw_event(GroupEventType.MESSAGE_EVENT, dataclass=MessageEvent)
async def handle_keyboard_events(event: MessageEvent):
    payload = event.object.payload
    user_id = event.object.user_id

    users = await bot.api.users.get(user_ids=[user_id])
    user = users[0]
    user_link = f"[id{user_id}|{user.first_name}{user.last_name}]"

    peer_id = event.object.peer_id

    logger.info("Пользователь %s нажал кнопку %s в чате %s",
                user_id, payload.get("act"), peer_id)
    
    if not payload:
        return
        
    # 1. Если нажали на стрелочку
    if payload.get("act") == "change_page":
        new_page = payload.get("page")
        method = payload.get("method")
        
        if method == "join_kb":
            # Редактируем сообщение: подменяем старую клавиатуру на новую
            await event.edit_message(
                message=f"Страница {new_page + 1}",
                keyboard=join_kb(page=new_page)
            )
        elif method == "hw_subj_kb":
             await event.edit_message(
                message=f"Страница {new_page + 1}",
                keyboard=hw_subj_kb(page=new_page, vk_id=peer_id)
            )
        
    # 2. Если нажали на само имя
    elif payload.get("act") == "select_name":
        name = payload.get("name")

        await bot.api.messages.delete(
            peer_id=event.object.peer_id,
            cmids=[event.object.conversation_message_id],
            delete_for_all=True
        )

        is_create = None

        if data.get_group_id(name=name) is not None:
            await event.show_snackbar("""Извините, но ваша заявка отклонена, т.к. такая группа уже привязана!
                                      (P.s. я сообщил администратору)""")
            is_create = False
            await bot.state_dispenser.delete(peer_id=peer_id)
        else:
            is_create = True
            await event.show_snackbar(f"Заявка на привязку вашего чата к группе {name} на рассмотрение!")
        
        admins = data.get_admins()
        for Id in admins:
            vk_id = data.get_user_vk_id(user_id=Id)

            if vk_id:
                if is_create is False:
                    await bot.api.messages.send(
                        user_id=vk_id,
                        random_id=randint(0,10000),
                        message=f"Пользователь {user_link} пытался присоеденить группу {name} к peer_id: {peer_id}"
                    )
                elif is_create is True:
                    await bot.api.messages.send(
                        user_id=vk_id,
                        random_id = randint(0,10000),
                        message = f"Пользователь {user_link} хочет создать группу {name} с peer_id: {peer_id}",
                        keyboard=confirm_adm_kb(name=name, peer_id=peer_id)
                    )
    elif payload.get("act") == "rejected":
        pl_peer_id = payload.get("peer_id")
        name = payload.get("name")
        
        await bot.api.messages.send(peer_id=pl_peer_id,
                                    message=f"Вам отказано в привязке вашего чата к группе {name}",
                                    random_id=randint(0,100000))
        await bot.api.messages.delete(
            peer_id = peer_id,
            delete_for_all=True,
            cmids=[event.object.conversation_message_id]
        )
        await bot.state_dispenser.delete(peer_id=pl_peer_id)

    elif payload.get("act") == "accepted":
        pl_peer_id = payload.get("peer_id")
        name = payload.get("name")

        if data.ensure_group(name=name, vk_id=pl_peer_id):
            await bot.api.messages.send(
                peer_id=pl_peer_id,
                random_id = randint(0,100000),
                message=f"Ваша заявка одобрена: чат был успешно привязан к группе {name}!"
            )
        else:
            await bot.api.messages.send(peer_id=pl_peer_id,
                                        random_id=randint(0,10000),
                                        message=f"Ваша чат по каким-то причинам не может быть привязана! Я уже сообщил Админу, подождите!")
            admins = data.get_admins()
            for Id in admins:
                vk_id = data.get_user_vk_id(user_id=Id)

                if vk_id:
                    await bot.api.messages.send(user_id=vk_id,
                                                random_id = randint(0,100000),
                                                message=f"Какие-то неполадки: не удалось привязать чат {pl_peer_id} к {name}")
            await bot.api.messages.delete(peer_id=peer_id,
                                          delete_for_all=True,cmids=[event.object.conversation_message_id])
        
        await bot.state_dispenser.delete(peer_id=pl_peer_id)
    
    elif payload.get("act") == "close":
        
        await bot.api.messages.delete(
            peer_id=event.object.peer_id,
            cmids=[event.object.conversation_message_id],
            delete_for_all=True
        )

        await event.show_snackbar("Меню закрыто")
    
    elif payload.get("act") == "get_hw":
        subj_id = payload.get("subj_id")
        subj_name = data.get_subject_name(subject_id=subj_id)

        await bot.api.messages.delete(
            peer_id=event.object.peer_id,
            cmids=[event.object.conversation_message_id],
            delete_for_all=True
        )

        hw = data.get_homework(vk_id=peer_id, subject_name=subj_name)
        lines = [f"ДЗ по {subj_name}:"]
        photos = []
        if hw is not None:
            for el in hw:
                if el[-2] in ([], None):
                    lines.append(f"-->{el[2]}")
                else:
                    lines.append(f"-->{el[2]} (Вложение...)")
                    photos.append(el[-2])
                
        else:
            lines.append("--> Не задано :|")
        
        text = "\n".join(lines)
        atts = []

        for collect in photos:
            for name in collect:
                file_path = ASSETS_DIR / name
                try:
                    att = await uploader.upload(str(file_path))
                    atts.append(att)
                except Exception:
                    logger.exception("Ошибка при выгрузке фото: ")

        keyboard = close_kb()
        await bot.api.messages.send(peer_id=peer_id, message=text, silent=True, attachment=atts, random_id = randint(0,100000), keyboard=keyboard)

    elif payload.get("act") == "del_hw":
        subj_id = payload.get("subj_id")
        subj_name = data.get_subject_name(subject_id=subj_id)
        hw = data.get_homework(vk_id = peer_id, subject_name = subj_name)
        if hw:
            if len(hw) == 1:
                hw_id = hw[0][0]
                data.delete_homework(hw_id)
                await bot.api.messages.edit(peer_id = peer_id, cmid = event.conversation_message_id,
                                        message = f"Дз для {subj_name} удалено")
                await asyncio.sleep(2)
                await bot.api.messages.delete(cmids=[event.conversation_message_id], peer_id = peer_id,
                                            delete_for_all=True)
            elif len(hw) > 1:
                kb = Keyboard(inline=True)
                for row in hw:
                    kb.add(
                        Callback(f"{row[2][:9]}...", payload={"act":"del_hw_2", "hw_id":row[0]}),
                        KeyboardButtonColor.PRIMARY
                    )
                    kb.row()
                    await bot.api.messages.edit(peer_id = peer_id, cmid = event.conversation_message_id,
                                        message = f"Выберите какое именно дз вы хотите удалить:", keyboard=kb)
        else:
            await bot.api.messages.edit(peer_id = peer_id, cmid = event.conversation_message_id,
                                        message = "Для этого предмета дз не найдено!")
            await asyncio.sleep(2)
            await bot.api.messages.delete(cmids=[event.conversation_message_id], peer_id = peer_id,
                                          delete_for_all=True)
    elif payload.get("act") == "del_hw_2":
        hw_id = payload.get("hw_id")
        data.delete_homework(hw_id)
        await bot.api.messages.edit(peer_id = peer_id, cmid = event.conversation_message_id,
                                        message = f"Дз удалено")
        await asyncio.sleep(2)
        await bot.api.messages.delete(cmids=[event.conversation_message_id], peer_id = peer_id,
                                    delete_for_all=True)


async def check_parser() -> None:
    logger.info("Проверка парсера началась")
    sch_date = data.get_last_schedule_date()
    filename = ASSETS_DIR / "last_date.pkl"
    
    with open(filename,'rb') as f:
        file_data = pickle.load(f)
    
    send_date = file_data.get("vk")
    
    if sch_date > send_date:
        file_data["vk"] = sch_date
        with open(filename,'wb') as f:
            pickle.dump(file_data, f)
        for name in data.get_group_names():
            schedule = data.get_schedule(group_name=name)
            vk_id = data.get_vk_id(group_name=name)
            now = datetime.strptime(sch_date, "%Y-%m-%d")
            days = ["ПОНЕДЕЛЬНИК", "ВТОРНИК", "СРЕДА", "ЧЕТВЕРГ", "ПЯТНИЦА", "СУББОТА", "ВОСКРЕСЕНЬЕ"]
            date = f"{days[now.weekday()]} ({now.strftime("%d.%m")})"
            lines = [f"Расписание на {date}", "---------------------"]
            ln = []
            subjects = []
            
            for lesson, subject, room in schedule:
                l = re.sub(r'\D', '', lesson)
                ln.append(l)
                s = subject.capitalize()
                subjects.append(s)
                r = "".join(c for c in room if c.isdigit()) if any(c.isdigit() for c in room) else room
                lines.append(f"({l})  {s}  [{r}]")

            lines.append("---------------------")
            photos = []

            for subj in subjects:
                hw = data.get_homework(vk_id=vk_id,subject_name=subj)
                if hw is not None:
                    lines.append(f"[{subj}]")
                    for el in hw:
                        lessons_left = el[-1] - 1
                        if lessons_left <= 0:
                            data.delete_homework(el[0])
                        else:
                            data.set_lesson(el[0],lessons_left=lessons_left)
                        
                        if el[-2] in ([], None):
                            lines.append(f"-->{el[2]}")
                        else:
                            lines.append(f"-->{el[2]} (Вложение...)")
                            photos.append(el[-2])
                atts = []

                for collect in photos:
                    for name in collect:
                        file_path = ASSETS_DIR / name
                        try:
                            att = await uploader.upload(str(file_path))
                            atts.append(att)
                        except Exception:
                            logger.exception("Ошибка при выгрузке фото: ")
                
            

            if min(ln) > "1":
                lines.append(f"*@all ВНИМАНИЕ! Завтра к {min(ln)} паре")
            
            text = "\n".join(lines).strip()
            logger.info("Бот отправил расписание группы %s в чат %s", name, vk_id)

            await bot.api.messages.send(
                peer_id=vk_id,
                random_id=randint(0,10000),
                message=text,
                attachment=atts
            )






async def periodic_task():
    while True:
        await check_parser()
        logger.info("Парсер проверен, я спать на 10 минут")
        await asyncio.sleep(600)



if __name__ == "__main__":
        try:
            bot.loop_wrapper.add_task(periodic_task())
            bot.run_forever()
        except KeyboardInterrupt:
            logger.info("Бот выключен! Спокойной ночи )")
        except Exception:
            logger.exception("Ошибка: ")