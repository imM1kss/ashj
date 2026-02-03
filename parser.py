from datetime import datetime, timedelta
import json
import os
import requests
import glob
from docx import Document
from vk_api import VkApi
from dotenv import load_dotenv

load_dotenv()

def run_parser():
    with open("data.json", "r", encoding="utf-8") as file:
        data = json.load(file)

    ndd, ndf = get_next_day()
    schedule = None

    if ndd != data.get("last_date"):
        download_link = get_vk_doc_link(data,ndd)

        if download_link:
            if download_file(download_link, ndf):
                file_name = f"{ndf}.docx"
                schedule = get_schedule(file_name)
                print(ndd)
                upload_data(data, "last_date", ndd)
                upload_data(data, "last_schedule", schedule)
    print(schedule)
    return schedule


def get_next_day():
    today = datetime.today()

    if today.weekday() == 5:  
        next_day = today + timedelta(days=2)
    elif today.weekday() == 6:
        next_day = today + timedelta(days=1)
    else:
        next_day = today + timedelta(days=1)

    ndd = next_day.strftime("%d.%m.%Y")
    ndf = next_day.strftime("%d_%m_%Y")

    return ndd, ndf

def get_vk_doc_link(data,ndd):
    print("ran get")
    vk_session = VkApi(token=os.getenv('access_token'))
    vk_api = vk_session.get_api()

    group_id = vk_api.groups.getById(group_id=data["url_schedule"])[0]["id"]

    wall = vk_api.wall.get(
        owner_id=-group_id,
        count=5
    )

    for post in wall.get("items", []):
        for att in post.get("attachments", []):
            print(att)
            if att["type"] == "doc" and att['doc']['title']== f"{ndd}.docx" :
                return att["doc"]["url"]

    return None

def download_file(download_link, ndf):
    if not download_link:
        return False

    filename = f"{ndf}.docx"
    file_path = os.path.join(os.getcwd(), filename)

    response = requests.get(download_link, stream=True)
    response.raise_for_status()

    remove_all_schedule()

    with open(file_path, "wb") as f:
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)

    return True

def get_schedule(doc_name):
    doc = Document(doc_name)
    schedule = []

    if not doc.tables:
        return schedule

    table = doc.tables[0]

    for row in table.rows:
        cells = clean([cell.text.strip() for cell in row.cells])

        if "25-14" in cells:
            cnt = len(cells)
            schedule.append([
                cells[cnt - 3],
                cells[cnt - 2],
                cells[cnt - 1]
            ])

    return schedule


def clean(items):
    result = []
    for item in items:
        if item and item not in result:
            result.append(item)
    return result

def remove_all_schedule():
    for file_path in glob.glob(os.path.join(os.getcwd(), "*.docx")):
        os.remove(file_path)

def upload_data(data, tag, value):
    print(f"Update {tag} to {value}")
    print(data[tag])
    data[tag] = value
    print(data[tag])
    with open("data.json", "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False, indent=2)

if __name__ == "__main__":
    result = run_parser()
    print(result)