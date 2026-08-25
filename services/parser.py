#imports
import os
import requests
import glob
import logging
import re
import json

from docx import Document
from vk_api import VkApi
from dotenv import load_dotenv
from datetime import datetime
from services.Datbase import DataBase
from typing import Optional, Dict, List
from time import sleep
from services.logging_config import setup_logging
from pathlib import Path

#logger init
setup_logging()
logger = logging.getLogger("parser")

#load values from .env
load_dotenv()

#const
TOKEN = os.getenv('access_token')
GROUP_ID = os.getenv('schedule_id')
ASSETS_DIR = Path('assets')


class VkGroup:
    global TOKEN, GROUP_ID #global values

    #init constructor
    def __init__(self, token=TOKEN, group_id=GROUP_ID):
        self.token = token
        self.group_id = group_id
    
    #connect to vk group schedule wall
    def _connect_wall(self) -> Dict:
        vk_session = VkApi(token=self.token) # session
        vk_api = vk_session.get_api() # connect api

        group_id = vk_api.groups.getById(group_id=self.group_id)[0]["id"] #get id group

        #get 5 last posts from wall
        wall = vk_api.wall.get(
            owner_id=-group_id,
            count=5
        )

        return wall
    
    #get date of last .docx with schedule
    def get_new_date(self) -> str:
        wall = self._connect_wall() #connection wall

        for post in wall.get("items", []): #iterate list with posts
            for att in post.get("attachments", []): #iterate attachemsts in every post

                att_type = att.get("type") # type of attachment

                if att_type == "doc": # only document type

                    att_title = att.get("doc", {}).get("title") #get title of document
                    date = text2date(att_title) #converting title to date

                    return date
        return ""
    
    #get file name of .docx with schedule
    def get_file_name(self) -> str:
        wall = self._connect_wall() #connecting wall

        for post in wall.get("items", []): #iterate list with posts
            for att in post.get("attachments", []):#iterate attachemsts in every post

                att_type = att.get("type") # type of attachment

                if att_type == "doc": # only document type

                    att_title = att.get("doc", {}).get("title") #get title of document

                    return att_title
        return ""


    # get schedule document downliad link
    def get_link(self) -> str:
        wall = self._connect_wall() #connecting wall

        file_name = self.get_file_name() #get file name

        for post in wall.get("items", []): #iterate list with posts
            for att in post.get("attachments", []): # iterate attacments in every post

                att_type = att.get("type") # type of attachment

                if att_type == "doc": #only document tipe

                    att_title = att.get("doc", {}).get("title") # doc title

                    if att_title == file_name: # only file with file_name

                        att_url = att.get("doc", {}).get("url") #get download link
                        logger.info("Parser get download link %s", att_url)

                        return att_url
        
        return ""


#main function
def run_parser() -> bool:
    try:
        logger.info("Парсер расписания запустился!")
        #Classes
        data = DataBase()
        group = VkGroup()
        #get dates
        last_date = data.get_last_schedule_date()
        last_date = last_date or ""
        logger.info(f"Парсер получил последнюю дату: {last_date}")
        new_date = group.get_new_date()
        logger.info(f"Парсер получил новую дату: {new_date}")


        if new_date > last_date:
            link = group.get_link()
            logger.info(f"Парсер получил ссылку на скачивание файла расписания: {link}")
            if download_file(link):
                logger.info(f"Парсер скачал файл расписания")
                schedule = get_schedule()
                logger.info(f"Парсер получил расписание")
                if schedule:
                    for cell in schedule:
                        if data.ensure_lesson(group_name=cell[0],
                                           subject_name=cell[2],
                                           lesson_num=cell[1],
                                           classroom=cell[3],
                                           date=new_date):
                            logger.info(f"Парсер занес в базу расписание для группы {cell[0]}")
                    return True
        return False
    except Exception:
        logger.exception("Exception:")

# function converter mixed names to numbers
def convert_group_name(group_name:Optional[str] = None) -> str:
    if group_name is None:
        return None
    
    result = re.sub(r'\D', '', group_name)
    return result 

def text2date(text:Optional[str] = None) -> str:
    if text is None:
        raise ValueError("Text is None")
    
    date_str = text.replace(".docx", "")
    dt = datetime.strptime(date_str, "%d.%m.%Y")
    date = dt.strftime("%Y-%m-%d")

    return date

def download_file(link:Optional[str] = None) -> bool:
    if link is None:
        return False

    #get file_name and file_path with schedule
    filename = ASSETS_DIR / "schedule.docx"
    file_path = os.path.join(os.getcwd(), filename)

    #request to link
    response = requests.get(link, stream=True)
    response.raise_for_status()

    logger.info("Parser take http request to %s. Status: %s", link, response.status_code)

    remove_all_schedule() # delete all *.docx

    #write schedule to schedule.docx
    with open(file_path, "wb") as f:
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)
    
    logger.info("File was downloaded. Path: %s", file_path)

    return True

#get schedule from schedule.docx
def get_schedule() -> List:
    doc = Document((ASSETS_DIR / "schedule.docx"))
    schedule = []

    if not doc.tables:
        return schedule

    table = doc.tables[0]
    data = DataBase()

    for row in table.rows:
        for name in data.get_group_names():
            cells = clean([cell.text.strip() for cell in row.cells])

            if len(cells) == 4:
                if name == convert_group_name(cells[0]):
                    schedule.append([
                        name,
                        cells[-3],
                        cells[-2],
                        cells[-1]
                    ])
            elif len(cells) > 4:
                if name == convert_group_name(cells[1]):
                    schedule.append([
                        name,
                        cells[-3],
                        cells[-2],
                        cells[-1]
                    ])
            elif len(cells) == 3:
                if name == convert_group_name(cells[0]):
                    schedule.append([
                        name,
                        cells[-2],
                        cells[-1],
                        "Пусто"
                    ])

    return schedule


def clean(items:List) -> List:
    result = []
    for item in items:
        if item and item not in result:
            result.append(item)
    return result

#delete all *.docx
def remove_all_schedule() -> None:
    for file_path in glob.glob(os.path.join(os.getcwd() / ASSETS_DIR, "*.docx")):
        os.remove(file_path)
    logger.info("All .docx removed")



if __name__ == '__main__':
    while True:
        try:
            result = run_parser()
            logger.info(f"Парсер выполнил свою работу и пошел спать! Результат - {result}")
            sleep(600)
        except KeyboardInterrupt:
            logger.info("Парсер выключен!")
            break
        except Exception:
            logger.info("Exception: ")
